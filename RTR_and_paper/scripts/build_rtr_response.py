#!/usr/bin/env python3
"""Build the final point-by-point RTR response as a polished Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAPER_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PAPER_ROOT / "reviewer" / "bioETH-PRS_RTR_response.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BLACK = RGBColor(0, 0, 0)


def set_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, 9, color=MUTED)


def set_reviewer_box(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.07

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    reviewer = doc.styles.add_style("Reviewer Comment", WD_STYLE_TYPE.PARAGRAPH)
    reviewer.base_style = normal
    reviewer.font.name = "Calibri"
    reviewer._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    reviewer._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    reviewer.font.size = Pt(10.5)
    reviewer.font.italic = True
    reviewer.font.color.rgb = GRAY
    reviewer.paragraph_format.left_indent = Inches(0.18)
    reviewer.paragraph_format.right_indent = Inches(0.08)
    reviewer.paragraph_format.space_before = Pt(4)
    reviewer.paragraph_format.space_after = Pt(8)
    reviewer.paragraph_format.line_spacing = 1.07

    meta = doc.styles.add_style("Response Metadata", WD_STYLE_TYPE.PARAGRAPH)
    meta.base_style = normal
    meta.font.name = "Calibri"
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    meta.font.size = Pt(9.5)
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.0


def add_metadata(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="Response Metadata")
    r = p.add_run(f"{label}: ")
    set_font(r, 9.5, bold=True, color=DARK_BLUE)
    r = p.add_run(value)
    set_font(r, 9.5, color=MUTED)


def add_comment_response(
    doc: Document,
    title: str,
    comment: str,
    locations: str,
    response: list[str],
) -> None:
    doc.add_heading(title, level=2)
    quote = doc.add_paragraph(comment, style="Reviewer Comment")
    set_reviewer_box(quote)
    add_metadata(doc, "Manuscript sections", locations)
    heading = doc.add_paragraph(style="Heading 3")
    heading.add_run("Response")
    for text in response:
        paragraph = doc.add_paragraph()
        for label in ("Previous wording", "Revised wording"):
            prefix = f"{label}:"
            if text.startswith(prefix):
                run = paragraph.add_run(prefix)
                set_font(run, 11, bold=True, color=DARK_BLUE)
                run = paragraph.add_run(text[len(prefix):])
                set_font(run, 11)
                break
        else:
            paragraph.add_run(text)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("RESPONSE TO REVIEWERS  |  bioETH-PRS")
    set_font(r, 9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Page ")
    set_font(r, 9, color=MUTED)
    add_page_field(fp)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Response to Reviewers")
    set_font(r, 23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("bioETH-PRS: Confidential Polygenic Risk Scoring with Smart Contracts on an FHE-Enabled Blockchain")
    set_font(r, 14, color=GRAY)
    add_metadata(doc, "Date", "1 August 2026")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "We thank the reviewers for their careful and constructive comments. The manuscript "
        "now states clearly what bioETH-PRS does, how it was evaluated, and where its limits "
        "remain. The reviewer comments are reproduced exactly below. Each response explains "
        "how the comment was addressed, identifies the relevant manuscript sections, and "
        "summarizes what the results show. Where a specialized term first matters, we begin "
        "with a short plain-language explanation before giving the technical details."
    )
    set_font(r, 11)

    doc.add_heading("Reviewer 1", level=1)
    add_comment_response(
        doc,
        "General assessment",
        "This manuscript presents bioETH-PRS, a privacy-preserving framework for polygenic risk score computation using fully homomorphic encryption on a programmable blockchain. The central idea is to replace the trusted evaluator used in prior encrypted PRS pipelines with auditable smart contracts, while protecting both patient genotypes and GWAS model weights. The manuscript is timely and conceptually interesting, particularly at the intersection of genomic privacy, encrypted computation, and decentralized infrastructure. However, the current evidence remains largely proof-of-concept, and several claims about deployability, privacy guarantees, and clinical feasibility are not yet fully supported.",
        "Abstract; Key Points; Introduction; Discussion; Conclusion",
        [
            "We thank the reviewer for identifying where the evidence and the claims were not yet aligned. We agree that the earlier version read too broadly for a proof-of-concept study. The revision now separates the architectural contribution from the empirical evidence: smart contracts coordinate the calculation without a designated evaluator, but the system still depends on contract correctness, blockchain consensus, and the fhEVM computation and decryption services.",
            "The empirical scope is now stated directly. We evaluated additive PRS models containing up to 5,000 variants. One public-weight 100-variant Classic calculation was completed on Sepolia. Streaming calculations with 100, 500, 1,000, and 5,000 variants, and the private-weight 100-variant calculation, were measured only in the local simulation. No local result is presented as a Sepolia, production, or clinical measurement.",
            "The revised manuscript also moves SNP-source authenticity into the main security discussion, replaces differential-privacy language with the more accurate term randomized risk category, expands the repeated-query stress test, and reports individual results for all 200 public-weight local comparisons. These changes make clear that the study does not establish clinical readiness, commercial practicality, genome-wide use, or a formal privacy guarantee.",
            "Previous wording: The abstract said that the approach “may be cost-competitive,” and the graphical abstract said that encrypted computation was “verified by blockchain consensus.”",
            "Revised wording: The Abstract, Key Points, Introduction, Discussion, and Conclusion now say that the study covers additive models up to 5,000 variants, distinguishes Sepolia from local measurements, and “reduces reliance on a single evaluator” while retaining the stated infrastructure dependencies.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 1 - Sepolia and local evaluation",
        "1. The empirical evaluation relies heavily on a mock coprocessor environment. The reported gas consumption, HCU budget, latency, and protocol behavior are mainly evaluated using a Hardhat in-process mock coprocessor rather than a real fhEVM deployment or public testnet. This substantially weakens the deployment claims. The authors should either provide real-network validation or clearly frame these results as simulation-based estimates.",
        "System design; Where calculations were evaluated; Transactions, gas, and fee examples; Discussion",
        [
            "A local simulation runs the contracts in a controlled development environment, whereas Sepolia is a public Ethereum test network. The distinction matters because a local result tests contract behavior but does not measure public-network delay, capacity, or cost.",
            "We thank the reviewer for asking us to draw this boundary more carefully. We completed one public-weight 100-SNP calculation on Sepolia using the Classic method (stored inputs). It required 25 transactions and 20,710,271 gas. The time from submission to the result was 269.3 seconds, followed by 8.1 seconds for decryption. The encoded score was 758,685, exactly matching the independent calculation.",
            "We identify local simulations separately throughout the manuscript. The same Classic method and 25-transaction arrangement used 18,755,864 gas in the local simulation; Sepolia therefore used 10.42% more gas in this comparison. This is a single comparison and is not used as a general conversion between local and public-network results.",
            "The public Streaming calculations at 100, 500, 1,000, and 5,000 variants were measured only in the local simulation. We did not obtain corresponding Sepolia Streaming measurements. The private-weight 100-SNP calculation was also evaluated only locally. We therefore report no Sepolia Streaming or private-weight time, gas, fee, capacity, or speed result.",
            "Previous wording: The evaluation described mock-coprocessor gas as “expected to be within 10–20% of real-network deployment” and presented local timings next to deployment-oriented claims.",
            "Revised wording: The sections “Where calculations were evaluated” and “Transactions, gas, and fee examples” label each result as Sepolia, local simulation, or calculated example, and state that local timing does not measure real encrypted computation, network latency, Sepolia capacity, or production fees.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 2 - Trust language",
        "2. The privacy claims should be stated more cautiously. The manuscript argues that bioETH-PRS removes the trusted evaluator assumption. This is a meaningful architectural contribution, but the system still depends on the correctness and availability of the fhEVM stack, smart contracts, ACL/decryption infrastructure, and blockchain consensus. Terms such as “zero trust” or “trustless” should be softened or carefully qualified.",
        "Title; Abstract; Key Points; Introduction; Security assumptions and limits; Discussion; Conclusion",
        [
            "A designated evaluator is one outside party that performs the encrypted calculation. bioETH-PRS removes that particular role, but it does not remove every dependency or source of trust.",
            "We thank the reviewer and agree with this distinction. The revised manuscript says that publicly auditable smart contracts reduce reliance on a single designated evaluator; it does not describe the complete system as zero-trust or trustless. The result still depends on correct genotype preparation, a valid model, the smart contracts, blockchain consensus, and the fhEVM calculation and decryption services.",
            "We also explain what consensus does and does not do. It orders and records contract execution according to the network rules. It does not independently prove the truth of the biological input, the clinical validity of the model, or the correctness of external encrypted-computation services.",
            "Previous wording: The graphical abstract said “Trustless” and “computation verified by blockchain consensus,” while the background called the blockchain “the trust anchor.”",
            "Revised wording: The Title, Abstract, Introduction, security section, Discussion, Conclusion, and graphical abstract now use “Consensus-enforced” or “reduces reliance on a single designated evaluator,” followed immediately by the remaining dependencies and limits.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 3 - Differential-privacy framing",
        "3. The noisy output oracle does not provide formal differential privacy. The authors acknowledge that the current mechanism is DP-inspired rather than a calibrated (epsilon, delta)-differential privacy guarantee. Given the sensitivity of genomic data, this limitation should be emphasized more prominently. If the authors wish to retain strong privacy language, they should provide a formal adjacency definition, sensitivity analysis, and privacy-parameter calibration.",
        "Randomized risk category; Limitations; Future work",
        [
            "Differential privacy is a formal mathematical guarantee about how much an output can reveal when one person's data change. Adding a random value is not enough by itself to establish that guarantee.",
            "We thank the reviewer for emphasizing this point. The manuscript now describes the feature as a randomized risk category and states explicitly that it does not provide differential privacy. A random integer from 0 through B-1 is added before the category is assigned. Its exact mean is (B-1)/2; in the local B=128 study, the mean is 63.5 and the contract uses an integer threshold adjustment of 64. We did not evaluate a randomized category release on Sepolia.",
            "For the 100-SNP data, the category remained unchanged for all 48 individuals whose scores were outside the uncertainty range. Two scores were inside the uncertainty range; one changed category in this calculation. A formal privacy guarantee would require a precise definition of which data sets are compared, how much one input can change the score, how the random value is chosen, and what repeated queries reveal.",
            "Previous wording: The manuscript called the component a “DP-inspired noisy output release mechanism” and described an “on-chain noisy output oracle,” wording that could be read as a stronger privacy claim than was demonstrated.",
            "Revised wording: The Abstract, architecture description, security section, limitations, conclusion, and figures use “randomized risk category (not differential privacy)” and explain the one-sided random addition, its bias, the uncertainty range, and the missing requirements for a formal guarantee.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 4 - Repeated-query analysis",
        "4. The manuscript estimates that model extraction would require thousands of hours under recommended rate-limiting settings. However, this calculation appears heuristic and does not fully address adaptive querying, multiple-wallet attacks, threshold manipulation, correlated SNP structure, or cross-sample probing. A stronger adversarial analysis is needed before the anti-probing claims can be considered established.",
        "Randomized risk category; Analysis of repeated queries; Limitations",
        [
            "Repeated-query analysis asks whether an authorized requester can learn a private model by submitting chosen genotype patterns and comparing the returned categories. One query is one complete calculation. An adaptive query is chosen after seeing earlier answers; a precommitted query is fixed before any answer is seen. This matters because adapting the next question can make extraction much more efficient.",
            "We thank the reviewer for pointing out that the earlier estimate was too abstract. We replaced the information-per-query heuristic with a controlled local stress test using a 20-weight private model. The number 20 was chosen so that each inferred coefficient could be checked separately; it is not a contract limit. An exact raw score is the baseline and revealed all 20 weights in 20 queries. Under a weaker interface in which the requester could change the threshold after each category, 19 of 20 estimates were within B after 200 queries, and all 20 were first within B after 260 queries.",
            "Here B=128 is the width of the integer random-addition range: the contract adds 0 through 127 before classification. “Within B” means an encoded-weight estimate has absolute error below 128 integer units; it is not exact recovery. Pearson correlation describes whether the estimated weights follow the true relative pattern, while sign accuracy says how often the positive or negative direction is correct. The 320-query rows use one common comparison budget; 320 is not a protocol security threshold.",
            "When requester-selected thresholds and inputs were all fixed in advance, none of the 20 estimates was within B after 320 queries; r was 0.6689 and 65% of signs were correct. Under the protocol's current provider-fixed threshold rule, none was within B after 320 queries, but r was 0.9388 and sign accuracy was 70%. The control therefore made close coefficient estimation harder in this example, but it did not completely hide relative information.",
            "Several wallets did not increase the quota for the same registered sample; separate registered samples had separate quotas. In a synthetic correlated-input test, each group of five SNPs was given the same dosage, reducing independent input variation, and r fell to 0.0223. This is not a biological assumption or a reliable defense because an authorized requester can submit other patterns. The fixed random sequence makes the example reproducible, not universal. At three calculations per 1,000 blocks, the 260-query adaptive example corresponds arithmetically to 288.9 hours at 12 seconds per block or 48.1 hours at 2 seconds per block; these are illustrative calculations, not measured delays or guarantees.",
            "Previous wording: The manuscript estimated “approximately 4,220 block windows” and “approximately 2,800 hours” from a bits-per-category heuristic without directly testing adaptive thresholds, wallet changes, sample changes, or correlated inputs.",
            "Revised wording: The section “Analysis of repeated queries” defines query, adaptive, precommitted, provider-fixed, B, Within B, Pearson r, sign accuracy, the 20-weight diagnostic model, the common 320-query budget, and the five-SNP correlated-block stress test before reporting the results and their limits.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 5 - SNP authenticity",
        "5. The inability to verify submitted encrypted SNPs is a major unresolved security issue. The system verifies access to a registered sample but cannot confirm that the submitted encrypted SNP values faithfully represent that sample. This allows malicious users to submit crafted inputs, which directly affects model-probing and misuse risks. This issue should be moved from a limitation to the main security discussion.",
        "Genotype preprocessing, QC, and model alignment; Security assumptions and limits; Limitations; Future work",
        [
            "Encryption can hide a submitted genotype value, but it does not prove where that value came from. Access permission and biological authenticity are therefore separate questions.",
            "We thank the reviewer and agree that this is a central security limitation, not a minor implementation detail. The main security discussion now states that the registry checks whether a requester may use a registered sample, but it cannot confirm that the encrypted SNP values came from that sample. A requester who may use the sample can therefore submit chosen values.",
            "The study assumes that the patient, laboratory, or data holder prepares the input correctly before encryption. The accompanying record identifies the genome build, variant order, and preparation rules, but it does not prove the biological origin of the values. Signed confirmation from a laboratory and privacy-preserving proof that an encrypted input matches a registered sample are described as future work.",
            "Previous wording: SNP provenance appeared only briefly in the limitations as “the system ... cannot verify that submitted encrypted SNP values faithfully represent that sample's genotype.”",
            "Revised wording: “Security assumptions and limits” now states in the main threat model that authorization does not establish biological source, and the preprocessing, limitations, and future-work sections identify who prepares the input and what additional proof would be needed.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 6 - Variant scale",
        "6. The prototype is evaluated on 100-5,000 SNP fixtures, whereas many PRS models contain tens of thousands to millions of variants. The authors should more clearly define the intended use case, such as curated small-panel PRS models, and avoid implying general applicability to large-scale clinical PRS deployment.",
        "Abstract; Key Points; Introduction; Variant scale; Discussion; Conclusion",
        [
            "The number of variants in a PRS varies widely. A calculation that works for a small or medium panel should not automatically be treated as evidence for a genome-wide score.",
            "We thank the reviewer and have narrowed the scope accordingly. The manuscript now states that this study evaluates additive PRS models containing up to 5,000 variants. The public-weight 100-SNP calculation was completed on Sepolia using the Classic method (stored inputs). In the local simulation, the Streaming method required 15, 47, 88, and 413 transactions for public-weight calculations with 100, 500, 1,000, and 5,000 variants, respectively. The private-weight 100-SNP calculation used the Streaming method and required 17 local transactions.",
            "The contracts process larger models in repeated groups and do not impose a fixed 100-variant model limit, but that architectural fact is not evidence of practical genome-wide scale. The 5,000-variant local calculation was the largest bioETH-PRS calculation evaluated. No Sepolia Streaming measurement was obtained, and the manuscript does not claim genome-wide or clinical PRS use.",
            "Previous wording: The comparison table described 5,000 variants as “scalable,” and the contribution list referred to “economically plausible” costs, which could imply broader deployment evidence.",
            "Revised wording: The Abstract, Key Points, intended-use paragraph, “Variant scale,” limitations, and Conclusion state “up to 5,000 variants,” identify which setting produced each result, and explicitly say that genome-wide and clinical use were not evaluated.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 7 - HEPRS comparison",
        "7. bioETH-PRS improves the trust model by removing the designated evaluator, but HEPRS supports much larger SNP counts and has different computational advantages. The manuscript should separate claims about privacy architecture, scalability, latency, memory use, and deployment assumptions rather than presenting bioETH-PRS as broadly superior.",
        "Comparison with HEPRS; Discussion; Conclusion",
        [
            "HEPRS and bioETH-PRS calculate the same type of weighted genetic score, but they place computation and trust in different parts of the system. Their evaluated sizes and performance measurements are therefore not directly interchangeable.",
            "We thank the reviewer and agree that the comparison should be dimension by dimension. The revised table reports what each system still depends on, its arithmetic method, evaluated variant count, timing, memory use, deployment requirements, released result, and publicly visible information. It does not describe either system as broadly superior.",
            "HEPRS reports encrypted computation with 110,000 variants and measured CKKS performance. bioETH-PRS reports a public-weight 100-SNP calculation on Sepolia using the Classic method (stored inputs) and uses smart contracts to record and control the calculation at a smaller scale. We do not compare local-simulation timing with HEPRS timing, and we state that memory use was not measured for bioETH-PRS.",
            "Previous wording: The comparison described bioETH-PRS as “evaluator-free,” listed “5,000 (scalable),” and compared HEPRS timing with a 386 ms local mock timing.",
            "Revised wording: “Comparison with HEPRS” and the Conclusion describe the systems as complementary, separate the trust architecture from scale and performance, report 110,000 variants for HEPRS versus 100 Classic on Sepolia and 100–5,000 Streaming locally for bioETH-PRS, and mark local timing as not comparable.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 8 - Cost claims",
        "8. The cost projections depend on L2-equivalent or application-chain gas pricing and are not based on measured production deployment. Claims that the system may be clinically or commercially practical should be toned down unless supported by real deployment data.",
        "Transactions, gas, and fee examples; Limitations; Conclusion",
        [
            "Gas is the accounting unit used for blockchain operations. Multiplying gas by an assumed gas price gives a fee example; it does not show what a production deployment would actually cost.",
            "We thank the reviewer and agree that the earlier pricing examples invited an unsupported practical interpretation. The cost section now reports measured Sepolia gas and test ETH without treating them as production prices or evidence of affordability. Deploying the four contracts on Sepolia used 5,892,559 gas across four transactions and 0.0062781714 Sepolia test ETH. The public-weight 100-SNP calculation used the Classic method (stored inputs), 20,710,271 gas, 25 transactions, and 0.0252747648 test ETH.",
            "In a separate local calculation using the Streaming method, the public-weight 100-SNP calculation used 15 transactions and 11.690 million gas, while the private-weight calculation used 17 transactions and 23.508 million gas, or 2.01 times as much. The fee examples multiply these local gas measurements by stated gas prices. They are calculations, not measured network costs or evidence of affordability.",
            "No corresponding Sepolia Streaming measurement was obtained, so the manuscript reports no Sepolia Streaming gas, time, fee, or affordability result.",
            "Previous wording: The manuscript projected USD prices and said the system “may become cost-competitive” and that a 5,000-SNP analysis could be “economically plausible.”",
            "Revised wording: “Transactions, gas, and fee examples,” limitations, and Conclusion report observed gas/test-ETH values separately from hypothetical fee calculations and state that production prices, USD cost, throughput, and commercial or clinical practicality were not measured.",
        ],
    )

    doc.add_heading("Reviewer 2", level=1)
    add_comment_response(
        doc,
        "General assessment",
        "This manuscript presents bioETH-PRS, a blockchain-based protocol for privacy-preserving polygenic risk scoring (PRS) using TFHE/fhEVM smart contracts. The paper’s main claim is that it removes the need for a trusted evaluator found in prior homomorphic-encryption PRS pipelines by moving orchestration to auditable on-chain contracts. Overall, the paper tries to addresses an important problem at the intersection of genomics, privacy, and decentralized computation. The manuscript is interesting and original. However, I do have several comments.",
        "Introduction; Genotype preprocessing, QC, and model alignment; Representing decimal weights as integers; Evaluation; Discussion; Conclusion",
        [
            "We thank the reviewer for the positive assessment and for the practical questions that follow. We revised the manuscript so that a reader can follow the calculation from a familiar PRS weighted sum through data preparation, encrypted integer encoding, contract execution, and decoding. The revision also separates what the protocol calculates from what must be supplied and validated by the laboratory, data preparer, model provider, and fhEVM services.",
            "The manuscript now explains genotype quality control and effect-allele alignment, includes a simple three-SNP worked example, identifies responsibility for each part of correctness, and compares every public-weight local result with an independent calculation of Equation 1. It reports all 200 person-by-model comparisons in a table, and it states that 5,000 variants was the largest model evaluated in the local simulation.",
            "Previous wording: The methods concentrated on contract and encoding details and did not provide a single plain-language path from genotype preparation to the decoded PRS or individual-level agreement for all participants.",
            "Revised wording: The Introduction, preprocessing/QC section, encoding and worked-example sections, evaluation, calculation-responsibility section, Discussion, and Conclusion now present that path and its evidence explicitly.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 1 - Practical variant scale",
        "1. bioETH-PRS was evaluated only on 100-5000 SNPs, while a real PRS in practice can involve far larger number of SNPs. Although the authors acknowledged that the HCU budget and transaction count made the genome-wide model impractical on current infrastructure. This is still a serious limitation because the method may only apply to a narrow class of PRS models with limited number of SNPs.",
        "Abstract; Introduction; Variant scale; Discussion; Limitations; Conclusion",
        [
            "A model's variant count is its number of genetic positions. More positions require more encrypted inputs and more blockchain transactions, even though the same calculation is repeated in groups.",
            "We thank the reviewer and agree that this is a serious scope limitation. We evaluated bioETH-PRS with additive PRS models containing up to 5,000 variants. A public-weight 100-SNP calculation was completed on Sepolia using the Classic method (stored inputs). Public-weight local calculations using the Streaming method required 15, 47, 88, and 413 transactions for 100, 500, 1,000, and 5,000 variants, respectively. The contracts process larger models in repeated groups and do not impose a fixed 100-variant model limit.",
            "The public-weight 5,000-variant local calculation was the largest bioETH-PRS calculation evaluated. No Sepolia Streaming measurement was obtained. We therefore frame the current use case as research on small-to-medium additive panels and state directly that genome-wide and clinical use were not demonstrated.",
            "Previous wording: The comparison table used “5,000 (scalable),” and the Discussion presented a practical-use split around the 5,000-SNP scale, which read as an application recommendation.",
            "Revised wording: The Abstract, Introduction, “Variant scale,” Discussion, limitations, and Conclusion state the tested maximum, transaction counts, evaluation setting, and the explicit non-claim for genome-wide or clinical use.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 2 - Genotype quality control",
        "2. Does bioETH-PRS require quality control of the genotype data, like missing value, minor allele frequency, etc? Please clarify this in the manuscript.",
        "Genotype preprocessing, QC, and model alignment",
        [
            "Genotype quality control means checking that the genetic values and their labels are usable before a score is calculated. Some checks concern one person, while others require a whole study cohort.",
            "We thank the reviewer for asking us to make this operational boundary explicit. The manuscript separates checks performed while a PRS model is developed from checks performed when one person is scored. Minor-allele-frequency and Hardy-Weinberg checks require a cohort and therefore occur before a model is published. bioETH-PRS scores one person at a time and cannot perform those cohort-level checks.",
            "Before encryption, the scoring data are checked for missing values, genome build, variant identity and order, allele orientation, duplicate variants, and unsupported variant types. Genotypes must be diploid values of 0, 1, or 2. Invalid values are rejected. The model must also state how missing values are handled; there is no unstated default. Build mismatches, duplicate or reordered variants, multiallelic sites, and insertions or deletions are rejected.",
            "Previous wording: The original manuscript did not contain a dedicated genotype-QC procedure and did not distinguish cohort-level model-development checks from single-person scoring checks.",
            "Revised wording: The new section “Genotype preprocessing, QC, and model alignment” states the pre-encryption checks, the rejection rules, the required missing-data policy, and why minor-allele-frequency and Hardy-Weinberg checks belong upstream.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 3 - Effect-allele alignment",
        "3. For some cases, the genotype of a SNP may be coded as 0, 1, 2 in terms of the number of risk alleles; but during the weights derivation, the genotype of that SNP in an independent dataset may be coded as 2, 1, 0 in terms of the number of minor alleles (when the risk allele is not the minor allele). Although we can require the genotype and the weights are provided with consistent coding, how to validate this requirement when they are totally blinded to each other? How does bioETH-PRS handle such situation?",
        "Introduction; Genotype preprocessing, QC, and model alignment",
        [
            "A dosage of 0, 1, or 2 counts copies of a chosen allele. The same genotype can appear reversed when one data set counts the effect allele and another counts the other allele.",
            "We thank the reviewer; this is exactly the kind of quiet coding mismatch that can reverse a score contribution. Equation 1 now defines each genotype value as the number of copies of the effect allele specified by the model, not the number of copies of the minor allele. The variant identity, genome build, effect allele, other allele, and model order are available for alignment even when the numerical weights are encrypted. Alignment is performed before the genotype values are encrypted.",
            "If the genotype already counts the effect allele, the value is kept. If it counts the other allele, the value is changed from g to 2-g. If the alleles match on the opposite DNA strand, the strand is corrected first. A/T and C/G variants whose strand cannot be resolved safely are rejected, as are incompatible and non-biallelic variants.",
            "Previous wording: Equation 1 described g only as “the allele dosage of a genetic variant,” without specifying which allele was counted or how reversal and strand cases were handled.",
            "Revised wording: The Introduction defines g as the model-specified effect-allele dosage, and the preprocessing section gives the keep, 2-g reversal, strand-correction, and rejection rules before encryption.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 4 - Who guarantees correctness?",
        "4. How and who to guarantee the final PRS provided by bioETH-PRS is correctly computed? In other words, the bioETH-PRS will eventually provide some numbers. But how do I know I can trust these numbers?",
        "Agreement with an independent calculation; Calculation checks and responsibilities",
        [
            "There are two different correctness questions: whether the arithmetic was carried out as specified, and whether the biological inputs and model were appropriate. The contract can address the first question for the submitted values, but not the second by itself.",
            "We thank the reviewer for asking the question in this direct way. The manuscript now explains who is responsible for each part of the result. The person or laboratory preparing the genotype data is responsible for alignment and quality checks. The model provider is responsible for the weights, thresholds, and scientific validity of the model. The contracts carry out the encoded weighted sum, while the fhEVM services perform the encrypted operations and apply the configured release permissions. Raw scores are authorized only for the requester, whereas randomized categories are publicly decryptable. An independent calculation of Equation 1 provides a numerical comparison.",
            "All 200 public-weight local calculations matched Equation 1 exactly. The public-weight Sepolia calculation also matched the independently calculated encoded score of 758,685. These results show that the calculation was correct for the studied inputs. They do not establish that the genotype values came from the biological sample, that the PRS model is clinically valid, or that it is accurate for populations not studied.",
            "Previous wording: The manuscript grouped “correctness and protocol verification” together and said that the expected 100-SNP encoded score was reproduced, without clearly separating arithmetic correctness from biological provenance and model validity.",
            "Revised wording: “Agreement with an independent calculation” reports all 200 comparisons, while “Calculation checks and responsibilities” names the data preparer, model provider, contracts, fhEVM services, and independent checker and states the limit of each role.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 5 - Explanation of the score calculation",
        "5. The original PRS calculation is simple and easy to understand/interpret, which is a weighted sum of multiple SNPs. The PRS calculation by bioETH-PRS seems more complicated with certain black boxes. Could the authors comment on that?",
        "Representing decimal weights as integers; Worked Example",
        [
            "Despite the extra encoding steps, bioETH-PRS still computes the familiar PRS weighted sum: each allele dosage is multiplied by its model weight and the products are added. The additional steps only let that calculation use unsigned encrypted integers.",
            "We thank the reviewer and agree that the mathematics should remain recognizable despite the encryption machinery. The explanation now starts with the familiar weighted sum in Equation 1. A three-SNP example gives a score of 0.45 by direct calculation. The manuscript then shows, in order, how the weights are converted to nonnegative integers, combined with the encrypted genotype values, corrected for the conversion, and converted back to the same score of 0.45.",
            "The manuscript now describes this order directly in the text: genotype checks and effect-allele alignment, encryption, contract calculation, and either requester-authorized raw-score release or public randomized-category release.",
            "Previous wording: The methods introduced a “three-step bijection” and several offsets before giving the reader a complete numerical example, making the ordinary weighted sum difficult to see.",
            "Revised wording: “Representing decimal weights as integers” first explains why unsigned encryption needs a shift, and “Worked Example” follows one three-SNP score from 0.45 through encoding and back to 0.45 in the same order used by the implementation.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 6 - Independent validation",
        "6. If I need double programming or independent validation of the final calculated PRS, could bioETH-PRS incorporate this?",
        "Where calculations were evaluated; Agreement with an independent calculation; Calculation checks and responsibilities",
        [
            "We thank the reviewer and agree that an independent implementation provides an important practical check on the calculation. To make this verification transparent, the manuscript includes a worked three-SNP example showing that the original PRS and its nonnegative integer representation produce the same score. We also independently recalculated Equation 1 for all 200 public-weight local test cases and obtained exact agreement with every bioETH-PRS result. The independent implementation and example inputs are included with the study materials so that others can repeat these checks. Although this agreement provides strong practical evidence that the algorithm has been implemented correctly, it is not a formal mathematical proof. Ethereum consensus orders and executes the deployed Solidity contract logic, but it does not independently recompute or verify the encrypted arithmetic. Correct encrypted arithmetic and decryption therefore still depend on the fhEVM services behaving as specified. These guarantees concern execution on the submitted encrypted values; they do not establish the biological correctness of those values. The network cannot determine whether the encrypted genotypes came from the stated individual, whether laboratory processing and quality control were performed correctly, or whether genome build, variant order, missing values, dosage coding, and effect-allele alignment were handled appropriately. These aspects remain the responsibility of the laboratory or data preparer, while the model provider remains responsible for the validity, calibration, and intended population of the PRS model.",
            "Previous wording: The evaluation reported that one expected 100-SNP encoded score was reproduced and that the two contract paths agreed, but it did not give an independent person-level comparison for every data set.",
            "Revised wording: “Agreement with an independent calculation” reports 50 people at each of four model sizes, the zero error summaries, 200/200 exact matches, and the reason exact agreement occurred for these particular weights.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 7 - Individual-level agreement",
        "7. In the Empirical Evaluation section, I was expecting to see that the individual PRS calculated by bioETH-PRS is consistent with the PRS calculated from Equation 1. Could the authors provide that information?",
        "Agreement with an independent calculation; Code Availability Statement",
        [
            "Individual-level agreement compares the two methods person by person, rather than reporting only an average error across the group.",
            "We thank the reviewer; we agree that the aggregate statement should be supported at the individual level. We calculated and decoded a separate public-weight result for each individual: 50 individuals for each of the 100-, 500-, 1,000-, and 5,000-variant data sets, for a total of 200. Every result matched the independent Equation 1 calculation exactly. The mean absolute error, root-mean-square error, and maximum absolute error were all zero, and the correlation was 1. The manuscript summarizes all 200 comparisons by model size in the agreement table.",
            "The manuscript summarizes these comparisons in the agreement table rather than a separate figure. The weights in these data sets have at most six decimal places, and the selected scale represented them without rounding. The exact agreement therefore applies to these data sets and does not establish accuracy for models with more precise weights. For the randomized risk category, all 48 individuals outside the uncertainty range kept the same category. Two individuals were inside the range; one changed category in this calculation.",
            "Previous wording: The manuscript said that reconstructed scores agreed “to machine epsilon” and highlighted one 100-SNP expected value, but it did not tabulate all 200 person-level comparisons.",
            "Revised wording: The agreement table gives the sample count and zero MAE/RMSE/maximum error for each model size, and the surrounding text states 200/200 exact matches, explains why no rounding occurred, and avoids a separate individual-agreement figure.",
        ],
    )

    doc.add_heading("Editor's Comments", level=1)
    add_comment_response(
        doc,
        "Editor",
        "(There are no comments.)",
        "Not applicable",
        ["We thank the editor. No separate editorial comments were provided."],
    )

    doc.core_properties.title = "Response to Reviewers - bioETH-PRS"
    doc.core_properties.subject = "Manuscript point-by-point response"
    doc.core_properties.author = "bioETH-PRS authors"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT.relative_to(PAPER_ROOT.parent)}")


if __name__ == "__main__":
    build()
