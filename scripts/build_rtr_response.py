#!/usr/bin/env python3
"""Build the final point-by-point RTR response as a polished Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "bioETH-PRS_RTR_response.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BLACK = RGBColor(0, 0, 0)


def set_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, 9, color=MUTED)


def set_reviewer_box(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    reviewer = doc.styles.add_style("Reviewer Comment", WD_STYLE_TYPE.PARAGRAPH)
    reviewer.base_style = normal
    reviewer.font.name = "Calibri"
    reviewer._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    reviewer._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    reviewer.font.size = Pt(10.5)
    reviewer.font.italic = True
    reviewer.font.color.rgb = GRAY
    reviewer.paragraph_format.left_indent = Inches(0.18)
    reviewer.paragraph_format.right_indent = Inches(0.08)
    reviewer.paragraph_format.space_before = Pt(4)
    reviewer.paragraph_format.space_after = Pt(8)
    reviewer.paragraph_format.line_spacing = 1.10

    meta = doc.styles.add_style("Response Metadata", WD_STYLE_TYPE.PARAGRAPH)
    meta.base_style = normal
    meta.font.name = "Calibri"
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    meta.font.size = Pt(9.5)
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.0


def add_metadata(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="Response Metadata")
    r = p.add_run(f"{label}: ")
    set_font(r, 9.5, bold=True, color=DARK_BLUE)
    r = p.add_run(value)
    set_font(r, 9.5, color=MUTED)


def add_comment_response(
    doc: Document,
    title: str,
    comment: str,
    actions: str,
    locations: str,
    evidence: str,
    response: list[str],
) -> None:
    doc.add_heading(title, level=2)
    quote = doc.add_paragraph(comment, style="Reviewer Comment")
    set_reviewer_box(quote)
    add_metadata(doc, "Completed actions", actions)
    add_metadata(doc, "Revised manuscript", locations)
    add_metadata(doc, "Evidence", evidence)
    heading = doc.add_paragraph(style="Heading 3")
    heading.add_run("Response")
    for text in response:
        doc.add_paragraph(text)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("RESPONSE TO REVIEWERS  |  bioETH-PRS revision")
    set_font(r, 9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Page ")
    set_font(r, 9, color=MUTED)
    add_page_field(fp)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Response to Reviewers")
    set_font(r, 23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("bioETH-PRS: Confidential Polygenic Risk Scoring with Auditable fhEVM Orchestration on a Programmable Blockchain")
    set_font(r, 14, color=GRAY)
    for label, value in (
        ("Manuscript status", "Revised submission"),
        ("Evidence boundary", "Stage A commit e6e2c1d; 174 automated tests passing"),
        ("Live validation", "Public 100-SNP Sepolia workflow completed; private weights mock-validated only"),
        ("Date", "31 July 2026"),
    ):
        add_metadata(doc, label, value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("We thank the reviewers for their careful and technically constructive assessment. We revised the implementation, evidence record, manuscript, and claims together. Reviewer text is reproduced verbatim below; every response identifies the completed action IDs, manuscript page/line locations in the line-numbered revision, and reproducible evidence.")
    set_font(r, 11)

    doc.add_heading("Reviewer 1", level=1)
    add_comment_response(
        doc,
        "General assessment",
        "This manuscript presents bioETH-PRS, a privacy-preserving framework for polygenic risk score computation using fully homomorphic encryption on a programmable blockchain. The central idea is to replace the trusted evaluator used in prior encrypted PRS pipelines with auditable smart contracts, while protecting both patient genotypes and GWAS model weights. The manuscript is timely and conceptually interesting, particularly at the intersection of genomic privacy, encrypted computation, and decentralized infrastructure. However, the current evidence remains largely proof-of-concept, and several claims about deployability, privacy guarantees, and clinical feasibility are not yet fully supported.",
        "All Reviewer 1 actions, R1.1 through R1.8",
        "Abstract and Key Points (p. 1, lines 1-48); Introduction (pp. 1-2, lines 49-152); Discussion and Conclusion (pp. 12-14, lines 838-1002)",
        "evidence/ phases 1-8; Stage A boundary e6e2c1d",
        [
            "We agree with the reviewer's characterization and have converted the paper from a broad deployment argument into a bounded, evidence-classed prototype study. The revision now separates Live fhEVM, Hardhat mock, and Analytic projection evidence; reports one verified public Sepolia workflow; defines retained trust and input-authenticity assumptions; replaces the incorrect anti-probing estimate with measured attacks; validates all 200 individual scores independently; and removes clinical or commercial practicality claims.",
            "The title and front matter now describe auditable fhEVM orchestration and evaluator minimization rather than trust removal. The revised conclusion states directly that bioETH-PRS is not a practical genome-wide PRS engine and that production affordability, clinical deployment, private-weight live execution, and a live HCU ceiling were not established.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 1 - Mock-only evaluation",
        "1. The empirical evaluation relies heavily on a mock coprocessor environment. The reported gas consumption, HCU budget, latency, and protocol behavior are mainly evaluated using a Hardhat in-process mock coprocessor rather than a real fhEVM deployment or public testnet. This substantially weakens the deployment claims. The authors should either provide real-network validation or clearly frame these results as simulation-based estimates.",
        "R1.1-E1, R1.1-E2, R1.1-M1",
        "Model Marketplace (p. 4, lines 287-301); Experimental Setup (pp. 9-10, lines 648-684); Tables 4-9 (pp. 10-12)",
        "evidence/phase7/live_2026-07-31/; evidence/phase8/scale_evidence.json; measured_transaction_use.json",
        [
            "We now provide real-network validation and label every result by evidence class. A public 100-SNP classic-path workflow completed on Sepolia (chain ID 11155111) with 25 status-1 workflow receipts, 20,710,271 gas, 269,320 ms submission-to-result latency, 8,081 ms Gateway/KMS decryption, and decoded encoded score 758,685, exactly matching the independent reference. Contract addresses, transaction hashes, block numbers, bytecode digests, source hash, and the reference digest are preserved in the receipt-level record.",
            "The first public attempt is retained as failed, not counted as success: nine transactions mined before a relayer TLS failure and no result was produced. We hardened the runner to prepare proofs before state-changing workflow writes, retry bounded transport failures, and checkpoint every receipt. Private-weight execution is implemented and Hardhat-mock validated, but the remaining 0.0127690815 Sepolia ETH did not safely cover the private workflow; no underfunded transaction was submitted. The manuscript therefore states private weights were not live-validated.",
            "A same-geometry public mock used 25 transactions and 18,755,864 gas, 10.42% below the live observation. We report that as one pair only and deleted the unsupported general 10-20% conversion claim. The live HCU ceiling remains unmeasured.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 2 - Trust language",
        "2. The privacy claims should be stated more cautiously. The manuscript argues that bioETH-PRS removes the trusted evaluator assumption. This is a meaningful architectural contribution, but the system still depends on the correctness and availability of the fhEVM stack, smart contracts, ACL/decryption infrastructure, and blockchain consensus. Terms such as “zero trust” or “trustless” should be softened or carefully qualified.",
        "R1.2-M1, R1.2-M2",
        "Title, Abstract, and Key Points (p. 1, lines 1-48); Introduction (p. 2, lines 90-101); Security Model and Table 2 (pp. 7-9, lines 492-584); Conclusion (p. 14, lines 972-1002)",
        "Table 2 trust/failure boundary; contracts and live provenance records",
        [
            "We agree. The revision consistently describes evaluator minimization: the designated application-level evaluator is replaced by publicly auditable contract orchestration, but trust is shifted rather than eliminated. The title was changed accordingly, and the abstract, graphical abstract, key points, introduction, comparison, discussion, and conclusion now use the same bounded vocabulary.",
            "A new trust and failure-boundary table names the genotype preprocessor, model provider, smart contracts, consensus, fhEVM coprocessor, Gateway/relayer, and ACL/threshold-decryption infrastructure and maps their failures to confidentiality, correctness, availability, and provenance. The manuscript also states that consensus makes contract execution auditable; it does not independently verify TFHE execution.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 3 - Differential-privacy framing",
        "3. The noisy output oracle does not provide formal differential privacy. The authors acknowledge that the current mechanism is DP-inspired rather than a calibrated (epsilon, delta)-differential privacy guarantee. Given the sensitivity of genomic data, this limitation should be emphasized more prominently. If the authors wish to retain strong privacy language, they should provide a formal adjacency definition, sensitivity analysis, and privacy-parameter calibration.",
        "R1.3-M1, R1.3-M2",
        "Bounded Randomized Categorical Release (pp. 8-9, lines 561-602); Limitations and Future Directions (p. 13, lines 891-924)",
        "contracts/ResultOracle.sol; test/rate_limit_randomized_release_test.ts; evidence/phase5/category_agreement_100snp.json",
        [
            "We removed the differential-privacy framing and use the implementation term 'bounded randomized categorical release.' The manuscript states explicitly that the mechanism provides no (epsilon, delta) guarantee because the noise is one-sided on [0,B), is not calibrated to score sensitivity, and has no composition accounting across repeated queries.",
            "The implemented behavior is described exactly: e_noisy = e + nu, model-defined thresholds fixed before any query, expected upward bias B/2, and the threshold-adjustment trade-off. In the measured 100-SNP classification study, both in-band individuals were exactly 64 = B/2 below their adjusted threshold, illustrating maximum uncertainty at a calibrated boundary. Formal adjacency, sensitivity calibration, signed noise, and composition analysis are now Future Directions rather than claimed properties.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 4 - Anti-probing evidence",
        "4. The manuscript estimates that model extraction would require thousands of hours under recommended rate-limiting settings. However, this calculation appears heuristic and does not fully address adaptive querying, multiple-wallet attacks, threshold manipulation, correlated SNP structure, or cross-sample probing. A stronger adversarial analysis is needed before the anti-probing claims can be considered established.",
        "R1.4-C1, R1.4-T1, R1.4-E1, R1.4-M1",
        "Fixed release policy (Algorithms 1-2, p. 7); Anti-Probing and Table 3 (pp. 9-10, lines 603-647); Limitations (p. 13, lines 874-897)",
        "evidence/phase6/anti_probing_results.json; scripts/anti_probing_evaluation.ts; contracts/attack-baseline/",
        [
            "We thank the reviewer for identifying a serious error. We deleted the submitted calculation rather than repairing it: it divided a count of candidate values by a bit rate, its stated 4,220 windows implied 14,067 rather than 2,800 hours, and the measured submitted-design cost is about 252 times lower than claimed.",
            "The replacement evaluation separates information cost from permitted rate. With rate limiting disabled, a raw-score path recovers 20/20 weights in 20 queries; the submitted caller-chosen-threshold interface recovers 20/20 in 200 adaptive queries; and the hardened fixed-threshold interface recovers 0/20 within B after 320 queries, while Pearson r = 0.9391 and 70% sign accuracy show residual model-shape leakage. At R=3 jobs per 1,000 blocks and an assumed 12-s block time, the measured submitted-interface cost is 11.1 hours per weight and 222.2 hours for the 20-weight model.",
            "Same-sample multi-wallet bypass is closed, but distinct samples retain independent quotas; private-model expansion additionally requires owner allowlisting. Correlated probes suppress the implemented estimator, but this is not an enforceable defense because arbitrary ciphertexts remain admissible. The revised claim is limited to reducing resolution and increasing measured query cost; it does not claim Sybil resistance or formal model confidentiality.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 5 - SNP authenticity",
        "5. The inability to verify submitted encrypted SNPs is a major unresolved security issue. The system verifies access to a registered sample but cannot confirm that the submitted encrypted SNP values faithfully represent that sample. This allows malicious users to submit crafted inputs, which directly affects model-probing and misuse risks. This issue should be moved from a limitation to the main security discussion.",
        "R1.5-T1, R1.5-M1, R1.5-M2",
        "Threat Model and Input Authenticity (pp. 7-8, lines 492-522); Table 2 (p. 9); Future Directions (p. 13, lines 913-918)",
        "test/prs_compute_engine_chunked_snp_test.ts trust-boundary regression; validation manifests",
        [
            "We agree and moved this boundary into the main Security Model. The threat model now includes a malicious authorized requester and states: contracts compute deterministically over submitted ciphertexts but do not prove those ciphertexts derive from the registered sample. GenomicRegistry.hasAccess gates who may create a job, not what is uploaded.",
            "The evaluated setting assumes trusted local preparation by the patient, an accredited laboratory, or an approved custodian. manifestHash is described as a provenance commitment to build, input hash, order, and preparation policy, not a cryptographic binding. The limitation is encoded in a renamed regression test that deliberately accepts arbitrary encrypted hard-call values after authorization. Signed laboratory attestation and zero-knowledge ciphertext-to-sample proofs are now explicit future work.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 6 - Bounded scale",
        "6. The prototype is evaluated on 100-5,000 SNP fixtures, whereas many PRS models contain tens of thousands to millions of variants. The authors should more clearly define the intended use case, such as curated small-panel PRS models, and avoid implying general applicability to large-scale clinical PRS deployment.",
        "R1.6-E1, R1.6-M1, R1.6-M2",
        "Introduction (p. 2, lines 148-152); Background (p. 3, lines 164-171); Bounded Scale Evidence and Table 6 (pp. 10-11, lines 712-725); Discussion (pp. 12-14, lines 838-881, 991-1002)",
        "evidence/phase8/scale_evidence.json; heprs_profile.json",
        [
            "We narrowed the intended use to a bounded-size research prototype for curated additive PRS models. One public 100-variant row is Live fhEVM. The wider executed range is Hardhat mock: public workflows at 100, 500, 1,000, and 5,000 variants required 15, 47, 88, and 413 host transactions; the private 100-variant mock workflow required 17.",
            "Rows at 10,000, 100,000, and 1,000,000 variants are labelled Analytic projection / unexecuted and carry no gas, latency, HCU, or fee claim. The abstract, key points, introduction, empirical evaluation, limitations, and conclusion all state that genome-wide execution and clinical deployment feasibility were not demonstrated.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 7 - HEPRS comparison",
        "7. bioETH-PRS improves the trust model by removing the designated evaluator, but HEPRS supports much larger SNP counts and has different computational advantages. The manuscript should separate claims about privacy architecture, scalability, latency, memory use, and deployment assumptions rather than presenting bioETH-PRS as broadly superior.",
        "R1.7-M1, R1.7-M2",
        "Comparison with HEPRS and Table 1 (pp. 4-5, lines 321-335); Complementary Systems (p. 13, lines 850-872); Conclusion (p. 14, lines 991-1002)",
        "Bundled HEPRS article, docs/PIIS2667237525003078.pdf; evidence/phase8/scale_evidence.json",
        [
            "We rebuilt the comparison one dimension per row: privacy architecture, designated evaluator, retained trust, arithmetic, demonstrated encrypted variants, latency evidence, memory evidence, deployment requirements, output policy, and metadata exposure. Every cell identifies its evidence type.",
            "The revision now states the trade-off directly. HEPRS demonstrates 110,000-variant real-FHE execution and measured CKKS performance, while bioETH-PRS demonstrates a public 100-SNP live fhEVM point and contract policy/auditability at smaller scale; bioETH-PRS memory was not measured. The in-process mock timing is no longer placed beside HEPRS real-FHE latency as if comparable, and all broad superiority language was removed.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 8 - Cost claims",
        "8. The cost projections depend on L2-equivalent or application-chain gas pricing and are not based on measured production deployment. Claims that the system may be clinically or commercially practical should be toned down unless supported by real deployment data.",
        "R1.8-E1, R1.8-M1",
        "Measured Transaction Use and Fee Sensitivity, Tables 4 and 9 (pp. 11-12, lines 744-774); Limitations (p. 13, lines 898-903); Conclusion (p. 14, lines 991-1002)",
        "evidence/phase8/measured_transaction_use.json; fee_sensitivity.json",
        [
            "We removed the projected USD table and all clinical/commercial affordability claims. The replacement separates observations from arithmetic. Live deployment used four transactions and 5,892,559 gas (0.0062781714 Sepolia test ETH); the live public job used 25 transactions and 20,710,271 gas (0.0252747648 test ETH). These are test-network expenditures, not production prices.",
            "Hardhat-mock public and private 100-variant streaming workflows used 15/17 transactions and 11.690/23.508 million gas; private is 2.01 times public and is the relevant model-extraction configuration. Hypothetical fee sensitivity is retained only as unexecuted ETH arithmetic, with no USD conversion or feasibility conclusion. The paper states production schedules, memory, throughput, and affordability were not measured.",
        ],
    )

    doc.add_heading("Reviewer 2", level=1)
    add_comment_response(
        doc,
        "General assessment",
        "This manuscript presents bioETH-PRS, a blockchain-based protocol for privacy-preserving polygenic risk scoring (PRS) using TFHE/fhEVM smart contracts. The paper’s main claim is that it removes the need for a trusted evaluator found in prior homomorphic-encryption PRS pipelines by moving orchestration to auditable on-chain contracts. Overall, the paper tries to addresses an important problem at the intersection of genomics, privacy, and decentralized computation. The manuscript is interesting and original. However, I do have several comments.",
        "Reviewer 2 actions R2.1 through R2.7",
        "Methods, Results, Discussion, and Conclusion (pp. 3-14)",
        "Independent validation and individual-level evidence under evidence/phase3 and evidence/phase5",
        [
            "We thank the reviewer. The revision preserves the architectural contribution while narrowing it to evaluator minimization under explicit infrastructure assumptions. We added executable preprocessing and alignment rules, a plain-language workflow, an independent implementation, a correctness-boundary table, and all 200 individual comparisons. We also surface the bounded-scale limitation throughout the paper.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 1 - Practical variant scale",
        "1. bioETH-PRS was evaluated only on 100-5000 SNPs, while a real PRS in practice can involve far larger number of SNPs. Although the authors acknowledged that the HCU budget and transaction count made the genome-wide model impractical on current infrastructure. This is still a serious limitation because the method may only apply to a narrow class of PRS models with limited number of SNPs.",
        "R2.1-M1; R1.6-E1, R1.6-M1, R1.6-M2",
        "Introduction (p. 2, lines 148-152); Bounded Scale Evidence (pp. 10-11, lines 712-725); Discussion and Limitations (pp. 12-13, lines 838-881)",
        "evidence/phase8/scale_evidence.json",
        [
            "We agree. The manuscript now calls bioETH-PRS a bounded-size research prototype for curated additive models and states that it is not a practical genome-wide PRS engine. The executed range and transaction counts are reported by evidence class, while the 10,000-1,000,000 rows are explicitly unexecuted transaction geometry. The limitation appears in the abstract, key points, introduction, opening of empirical evaluation, discussion, and conclusion rather than only in retrospective limitations.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 2 - Genotype quality control",
        "2. Does bioETH-PRS require quality control of the genotype data, like missing value, minor allele frequency, etc? Please clarify this in the manuscript.",
        "R2.2-C1, R2.2-T1, R2.2-M1",
        "Genotype Preprocessing, QC, and Model Alignment (p. 3, lines 172-217)",
        "validation/independent_prs_reference.py; validation/cases/; 56/56 reference self-tests",
        [
            "We added a Methods subsection transcribed from the independent validator. MAF and Hardy-Weinberg filtering are cohort/model-development QC performed upstream because bioETH-PRS scores one individual and never observes a cohort. Scoring-time checks cover missingness, genome build, variant identity/order, allele orientation, duplicates, representation, and site type.",
            "Only integer diploid hard calls in {0,1,2} are accepted; invalid values are rejected, never clamped. Missingness policy is a required manifest field (reject, zero_dosage, or mean_dosage) with no default. Build mismatch, reordered or duplicate variants, multiallelic sites, and indels are rejected. Every run reports matched, intercept, missing, imputed, invalid, and rejected counts. The bare HEPRS matrices lack metadata and are explicitly treated as pre-aligned fixtures.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 3 - Effect-allele alignment",
        "3. For some cases, the genotype of a SNP may be coded as 0, 1, 2 in terms of the number of risk alleles; but during the weights derivation, the genotype of that SNP in an independent dataset may be coded as 2, 1, 0 in terms of the number of minor alleles (when the risk allele is not the minor allele). Although we can require the genotype and the weights are provided with consistent coding, how to validate this requirement when they are totally blinded to each other? How does bioETH-PRS handle such situation?",
        "R2.3-C1, R2.3-T1, R2.3-M1",
        "Equation 1 definition (p. 1, lines 50-55); Effect-allele alignment (p. 3, lines 196-217)",
        "validation/independent_prs_reference.py harmonize_dosage; allele-orientation known-answer tests",
        [
            "We now define g_i as the dosage of the model-specified effect allele, not the minor-allele count. Alignment does not require revealing weight magnitudes: public model metadata exposes variant identity, genome build, effect allele, other allele, and order even when the weights are encrypted, and alignment occurs locally before encryption.",
            "If the genotype already counts the effect allele, g is retained; if it counts the other allele, the validator applies 2-g. Compatible strand complements are resolved before the same decision. Unresolved palindromic A/T and C/G sites are rejected because a literal label match is strand-ambiguous; incompatible and non-biallelic sites are also rejected. Match, flip, ambiguity, and rejection counts are emitted and covered by hand-checked tests.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 4 - Who guarantees correctness?",
        "4. How and who to guarantee the final PRS provided by bioETH-PRS is correctly computed? In other words, the bioETH-PRS will eventually provide some numbers. But how do I know I can trust these numbers?",
        "R2.4-E1, R2.4-M1",
        "Quantisation Accuracy and Figure 7 (p. 10, lines 685-711); Correctness and Protocol Verification (p. 12, lines 775-786); Table 10 (p. 13)",
        "evidence/phase5/individual_level_comparison.csv; provenance blocks; on-chain receipt verification",
        [
            "We replaced an undifferentiated correctness claim with a boundary table. The genotype preprocessor guarantees alignment/QC rules; the model provider is responsible for weights, thresholds, and scientific validity; contracts deterministically implement the encoded weighted sum; fhEVM infrastructure is responsible for encrypted execution and authorized decryption; the independent implementation checks agreement with Equation 1; and the end user can verify manifests, addresses, bytecode digests, receipts, and the reference digest.",
            "All 200 independently compared contract jobs agree exactly with Equation 1, and the public live job matches encoded score 758,685. These results validate the tested pipeline but do not guarantee biological sample authenticity, model validity, calibration, ancestry portability, or untested deployments; the manuscript states each exclusion explicitly.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 5 - Interpretability of the encoded pipeline",
        "5. The original PRS calculation is simple and easy to understand/interpret, which is a weighted sum of multiple SNPs. The PRS calculation by bioETH-PRS seems more complicated with certain black boxes. Could the authors comment on that?",
        "R2.5-M1",
        "Three-Step Unsigned Encoding and worked example (pp. 5-6, lines 349-387); Figure 4 (p. 6); Algorithms 1-2 (p. 7)",
        "validation/cases/mixed_signed_weights.json; independent reference self-test",
        [
            "We reorganized the explanation around the original weighted sum. The three-SNP example now begins with the plaintext result 0.45, then shows six reversible steps: quantize, shift weights, accumulate, correct the weight zero-point, shift the score, and decode. The decoded result returns exactly to Equation 1.",
            "A new six-stage workflow figure separates metadata/QC and effect-allele orientation from encryption, contract accumulation, ACL-gated release, and decoding. Handle details remain in the protocol section rather than obscuring the arithmetic. The worked example is also a hand-checked independent validation case.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 6 - Independent validation",
        "6. If I need double programming or independent validation of the final calculated PRS, could bioETH-PRS incorporate this?",
        "R2.6-C1, R2.6-T1",
        "Experimental Setup and Reproducibility Identifiers (pp. 9-10, lines 648-684); Correctness and Protocol Verification (p. 12, lines 775-786)",
        "validation/independent_prs_reference.py; npm run validate:cross-language; evidence/phase3/",
        [
            "Yes. We added a standard-library-only Python reference implementation derived independently from the TypeScript/Solidity path. One command runs its 56-rule self-test, scores three hand-computed cases, executes the TypeScript/contract path, and compares encoded and decoded outputs at tolerance zero. All three cases pass with zero mismatches.",
            "Every evidence-producing runner now records the model/fixture/manifest hashes, contract addresses and bytecode digests, repository/source identity, and independent reference-output digest. This supports repeatable double programming without claiming that agreement between two implementations is a formal proof.",
        ],
    )

    add_comment_response(
        doc,
        "Comment 7 - Individual-level agreement",
        "7. In the Empirical Evaluation section, I was expecting to see that the individual PRS calculated by bioETH-PRS is consistent with the PRS calculated from Equation 1. Could the authors provide that information?",
        "R2.7-E1, R2.7-M1",
        "Quantisation Accuracy, Table 5, and Figure 7 (pp. 10-11, lines 685-711); supplementary 200-row CSV",
        "evidence/phase5/individual_level_comparison.csv; summary_statistics.json; category_agreement_100snp.json",
        [
            "We have now executed and decoded a separate contract job for every individual: 50 individuals at each of 100, 500, 1,000, and 5,000 nominal variants, 200 jobs total. Against the independent Equation 1 implementation, MAE, RMSE, and maximum absolute error are all zero, exact matches are 200/200, and Pearson r = 1 in exact decimal arithmetic. The scatter plot places every point on the identity line, and all 200 rows are supplied.",
            "We also clarify why zero error does not generalize: all 6,604 fixture weights have at most six decimals and the selected scale is an integer multiple of 10^6, so quantization is lossless and no rounding occurs. The result validates the end-to-end tested pipeline, not universal precision. Category agreement is reported as 48/48 outside the noise band, with two in-band individuals listed separately rather than folded into a misleading 50/50 claim.",
        ],
    )

    doc.add_heading("Editor's Comments", level=1)
    add_comment_response(
        doc,
        "Editor",
        "(There are no comments.)",
        "No action required",
        "Not applicable",
        "Not applicable",
        ["We thank the editor. No separate editorial comments required a response."],
    )

    doc.add_heading("Revision Verification Summary", level=1)
    for text in (
        "Manuscript compile: successful line-numbered 15-page PDF with resolved bibliography.",
        "Automated suite: 174 passing, 0 failing on Node v22.23.1.",
        "Independent validation: three known-answer cases pass at tolerance zero; 56 reference self-checks pass.",
        "Live verification: four Sepolia deployments and all 25 successful public-workflow receipts independently rechecked; decoded encoded score 758,685.",
        "Evidence boundary: all manuscript measurements originate in evidence/ and are labelled Live fhEVM, Hardhat mock, or Analytic projection.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        p.add_run(text)

    doc.core_properties.title = "Response to Reviewers - bioETH-PRS"
    doc.core_properties.subject = "Revised manuscript point-by-point response"
    doc.core_properties.author = "bioETH-PRS authors"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    build()
